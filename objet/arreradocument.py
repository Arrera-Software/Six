from docx import Document
from odf.opendocument import OpenDocumentText
from odf import text, opendocument ,teletype
from odf.opendocument import load
from odf.text import P

class CArreraDocx :
    def __init__(self,fileName:str) :
        if "docx" in fileName:
            self.__document = Document()
            self.__fileName = fileName
            self.__typeFile = "docx"
            self.__init = True
        elif "odt"in fileName:
            self.__document = opendocument.OpenDocumentText()
            self.__fileName = fileName
            self.__typeFile = "odt"
            self.__init = True
        else :
            self.__fileName = fileName
            self.__init = False
    
    def writeNoEcrase(self,content:str):
        if self.__init:
            if self.__typeFile== "docx":
                try :
                    for line in content.split('\n'):
                        self.__document.add_paragraph(line)
                    self.__document.save(self.__fileName)
                    return True
                except Exception as e:
                    # print(f"Error writing to docx file: {e}")
                    return False
            elif self.__typeFile== "odt":
                try :
                    for line in content.split('\n'):
                        p = text.P()
                        p.addText(line)
                        self.__document.text.addElement(p)
                    self.__document.save(self.__fileName)
                    return True
                except Exception as e:
                    # print(f"Error writing to odt file: {e}")
                    return False
            else :
                return False
        else :
            return False
    
    def writeEcrase(self,content:str):
        if self.__init:
            if self.__typeFile== "docx":
                try :
                    del self.__document
                    self.__document = None
                    self.__document = Document()
                    self.__document.add_paragraph(content)
                    self.__document.save(self.__fileName)
                    return True
                except Exception as e:
                    # print(f"Error writing to docx file: {e}")
                    return False
            elif self.__typeFile== "odt":
                try :
                    del self.__document
                    self.__document = None
                    self.__document = opendocument.OpenDocumentText()
                    paragraphe = P(text=content)
                    self.__document.text.addElement(paragraphe)
                    self.__document.save(self.__fileName)
                    return True
                except Exception as e:
                    # print(f"Error writing to odt file: {e}")
                    return False
            else :
                return False
        else :
            return False
    
    def read(self):
        if self.__init:
            if self.__typeFile== "docx":
                self.__document = Document(self.__fileName)
                texte = []
                for paragraph in self.__document.paragraphs:
                    texte.append(paragraph.text)
                return '\n'.join(texte)
            elif self.__typeFile== "odt":
                    self.__document = load(self.__fileName)
                    content = ""
                    for para in self.__document.getElementsByType(text.P):
                        content += teletype.extractText(para)+"\n"
                    return content
            else :
                return "error"
        else :
            return "error"